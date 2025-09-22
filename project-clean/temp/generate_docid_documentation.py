#!/usr/bin/env python3
"""
Generate comprehensive DOCX documentation from documentation.md
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("python-docx not installed. Installing...")
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

import re
import datetime
import markdown
from markdown.extensions import codehilite

def add_image_placeholder(doc, text, width=6, height=3):
    """Add a placeholder for an image with descriptive text"""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add a bordered table as image placeholder
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.width = Inches(width)
    
    # Add placeholder text
    cell_paragraph = cell.paragraphs[0]
    cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_paragraph.add_run(f"[INSERT SCREENSHOT/DIAGRAM]\n{text}")
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Add some spacing
    doc.add_paragraph()

def setup_styles(doc):
    """Set up custom styles for the document"""
    
    # Create custom Code style if it doesn't exist
    try:
        code_style = doc.styles['Code']
    except KeyError:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'Consolas'
        code_font.size = Pt(9)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        code_style.paragraph_format.keep_together = True
    
    # Create custom Command style
    try:
        command_style = doc.styles['Command']
    except KeyError:
        command_style = doc.styles.add_style('Command', WD_STYLE_TYPE.PARAGRAPH)
        command_font = command_style.font
        command_font.name = 'Courier New'
        command_font.size = Pt(8)
        command_font.bold = True
        command_style.paragraph_format.left_indent = Inches(0.3)
        command_style.paragraph_format.space_before = Pt(3)
        command_style.paragraph_format.space_after = Pt(3)

def parse_markdown_content(file_path):
    """Parse the markdown content from documentation.md"""
    
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()
    
    # Split content into sections
    sections = []
    current_section = {"title": "", "content": ""}
    
    lines = content.split('\n')
    
    for line in lines:
        # Check for headers
        if line.startswith('# '):
            if current_section["title"]:
                sections.append(current_section)
            current_section = {"title": line[2:], "level": 1, "content": ""}
        elif line.startswith('## '):
            if current_section["title"]:
                sections.append(current_section)
            current_section = {"title": line[3:], "level": 2, "content": ""}
        elif line.startswith('### '):
            if current_section["title"]:
                sections.append(current_section)
            current_section = {"title": line[4:], "level": 3, "content": ""}
        elif line.startswith('#### '):
            if current_section["title"]:
                sections.append(current_section)
            current_section = {"title": line[5:], "level": 4, "content": ""}
        else:
            current_section["content"] += line + '\n'
    
    # Add the last section
    if current_section["title"]:
        sections.append(current_section)
    
    return sections

def add_content_to_doc(doc, sections):
    """Add parsed content to the document"""
    
    # Title page
    title = doc.add_heading('DOCiD Flask API', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Comprehensive Technical Documentation', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Document info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(f"Generated on: {datetime.datetime.now().strftime('%B %d, %Y')}")
    info_run.font.size = Pt(12)
    
    doc.add_paragraph()
    add_image_placeholder(doc, "System Architecture Overview", 8, 5)
    
    # Page break
    doc.add_page_break()
    
    # Process each section
    for section in sections:
        if not section["title"] or not section["content"].strip():
            continue
            
        # Add heading based on level
        level = section.get("level", 1)
        if level <= 3:
            doc.add_heading(section["title"], level=level)
        else:
            para = doc.add_paragraph()
            run = para.add_run(section["title"])
            run.bold = True
            run.font.size = Pt(11)
        
        # Process content
        content = section["content"].strip()
        
        # Split content by code blocks and regular text
        parts = re.split(r'```(\w*\n.*?\n)```', content, flags=re.DOTALL)
        
        for i, part in enumerate(parts):
            if i % 2 == 0:  # Regular text
                # Process regular paragraphs
                paragraphs = part.split('\n\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if not para_text:
                        continue
                    
                    # Handle lists
                    if para_text.startswith('- ') or para_text.startswith('* '):
                        lines = para_text.split('\n')
                        for line in lines:
                            if line.strip().startswith(('- ', '* ')):
                                doc.add_paragraph(line.strip()[2:], style='List Bullet')
                            elif line.strip().startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
                                doc.add_paragraph(line.strip()[3:], style='List Number')
                    
                    # Handle bold text patterns
                    elif '**' in para_text:
                        para = doc.add_paragraph()
                        
                        # Split by bold markers
                        bold_parts = re.split(r'\*\*(.*?)\*\*', para_text)
                        for j, bold_part in enumerate(bold_parts):
                            if j % 2 == 0:  # Regular text
                                if bold_part.strip():
                                    para.add_run(bold_part)
                            else:  # Bold text
                                if bold_part.strip():
                                    run = para.add_run(bold_part)
                                    run.bold = True
                    
                    # Regular paragraph
                    elif para_text:
                        doc.add_paragraph(para_text)
                        
                        # Add image placeholders for key sections
                        if any(keyword in section["title"].lower() for keyword in 
                               ['architecture', 'database', 'deployment', 'monitoring', 'api', 'swagger', 'postman']):
                            if any(trigger in para_text.lower() for trigger in 
                                   ['diagram', 'interface', 'screenshot', 'dashboard', 'configuration']):
                                add_image_placeholder(doc, f"{section['title']} - {para_text[:50]}...")
            
            else:  # Code block
                code_text = part.strip()
                if code_text:
                    # Remove language identifier from first line
                    code_lines = code_text.split('\n')
                    if code_lines[0].strip() in ['bash', 'python', 'json', 'http', 'javascript', 'nginx', 'sql']:
                        code_text = '\n'.join(code_lines[1:])
                    
                    # Add code block
                    if len(code_text) > 500:  # Long code blocks
                        # Split into smaller chunks
                        chunks = [code_text[i:i+500] for i in range(0, len(code_text), 500)]
                        for chunk in chunks:
                            doc.add_paragraph(chunk, style='Code')
                    else:
                        doc.add_paragraph(code_text, style='Code')

def create_table_of_contents(doc, sections):
    """Create a table of contents"""
    
    doc.add_heading('Table of Contents', level=1)
    
    # Extract main sections
    toc_items = []
    section_counter = 1
    
    for section in sections:
        if section.get("level", 1) == 1 and section["title"].strip():
            toc_items.append(f"{section_counter}. {section['title']}")
            section_counter += 1
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()

def main():
    """Main function to generate the DOCX documentation"""
    
    print("🔄 Generating DOCiD Documentation from documentation.md...")
    
    # Create document
    doc = Document()
    
    # Set up styles
    setup_styles(doc)
    
    # Parse markdown content
    documentation_path = "/Users/ekariz/Projects/AMBAND/DOCiD/backend/documentation.md"
    sections = parse_markdown_content(documentation_path)
    
    print(f"📖 Parsed {len(sections)} sections from documentation.md")
    
    # Add content to document
    add_content_to_doc(doc, sections)
    
    # Add specific sections that might need special handling
    special_sections = [
        {
            "title": "Port Configuration Reference",
            "content": """
Production Port Configuration:
• Flask Production: Port 6000
• Flask Development: Port 5001  
• Node.js Production: Port 6001
• Node.js Development: Port 5002

Note: Port 5000 conflicts with macOS AirPlay service.
"""
        },
        {
            "title": "Quick Command Reference", 
            "content": """
Essential Commands:
• Start Development: flask run --port=5001 --debug
• Production Deploy: gunicorn --config gunicorn.conf.py wsgi:app
• Supervisor Control: sudo supervisorctl restart docid
• Database Backup: pg_dump -U usr_docid -d docid | gzip > backup.sql.gz
• View Logs: tail -f logs/gunicorn.out.log
• Push to CORDRA: python push_to_cordra.py --publication-id ID
"""
        }
    ]
    
    for special_section in special_sections:
        doc.add_heading(special_section["title"], level=2)
        doc.add_paragraph(special_section["content"])
        add_image_placeholder(doc, f"{special_section['title']} Screenshot", 6, 3)
    
    # Save document
    output_filename = "/Users/ekariz/Projects/AMBAND/DOCiD/backend/docid-documentation.docx"
    doc.save(output_filename)
    
    print(f"✅ Documentation generated successfully: {output_filename}")
    
    # Get file size
    import os
    file_size = os.path.getsize(output_filename)
    file_size_mb = file_size / (1024 * 1024)
    
    print(f"📊 Document Statistics:")
    print(f"   • File size: {file_size_mb:.2f} MB")
    print(f"   • Sections processed: {len(sections)}")
    print(f"   • Content includes:")
    print(f"     - Complete deployment guide (10 detailed sections)")
    print(f"     - API documentation with Swagger and Postman examples") 
    print(f"     - CORDRA integration with recent enhancements")
    print(f"     - Production configuration and monitoring")
    print(f"     - Troubleshooting and performance optimization")
    print(f"     - Image placeholders for screenshots and diagrams")
    
    return output_filename

if __name__ == "__main__":
    try:
        filename = main()
        print(f"\n🎉 Success! Open the file: {filename}")
        
    except FileNotFoundError:
        print("❌ Error: documentation.md file not found!")
        print("Make sure the file exists at: /Users/ekariz/Projects/AMBAND/DOCiD/backend/documentation.md")
        
    except Exception as e:
        print(f"❌ Error generating documentation: {str(e)}")
        import traceback
        traceback.print_exc()