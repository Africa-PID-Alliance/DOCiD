#!/usr/bin/env python3
"""
Script to generate comprehensive DOCX documentation for DOCiD Flask API
with image placeholders for screenshots
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("python-docx not installed. Installing...")
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH

import datetime

def add_image_placeholder(doc, text, width=6, height=4):
    """Add a placeholder for an image with descriptive text"""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add a bordered table as image placeholder
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.width = Inches(width)
    
    # Add placeholder text
    cell_paragraph = cell.paragraphs[0]
    cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_paragraph.add_run(f"[INSERT SCREENSHOT HERE]\n{text}")
    run.font.size = Pt(12)
    run.italic = True
    
    # Add some spacing
    doc.add_paragraph()

def create_docx_documentation():
    """Create comprehensive DOCX documentation"""
    
    # Create document
    doc = Document()
    
    # Set up styles
    title_style = doc.styles['Title']
    heading1_style = doc.styles['Heading 1']
    heading2_style = doc.styles['Heading 2']
    heading3_style = doc.styles['Heading 3']
    
    # Create custom Code style if it doesn't exist
    try:
        code_style = doc.styles['Code']
    except KeyError:
        # Create a custom code style
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'Courier New'
        code_font.size = Pt(9)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
    
    # Title Page
    title = doc.add_heading('DOCiD Flask API', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Comprehensive Technical Documentation', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Document info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(f"Generated on: {datetime.datetime.now().strftime('%B %d, %Y')}")
    info_run.font.size = Pt(12)
    
    doc.add_paragraph()
    add_image_placeholder(doc, "System Architecture Overview Diagram", 7, 5)
    
    # Page break
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        "1. Executive Summary",
        "2. System Architecture", 
        "3. Installation & Setup",
        "4. Database Schema & Data Models",
        "5. API Documentation",
        "6. External Service Integrations",
        "7. Comments System",
        "8. Authentication & Security",
        "9. Deployment Guide",
        "10. Testing & Quality Assurance",
        "11. Monitoring & Maintenance",
        "12. Appendices"
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    
    doc.add_paragraph(
        "The DOCiD (Document ID) system is a sophisticated Flask-based REST API designed to serve as a "
        "comprehensive publication and document identifier management platform. The system provides "
        "persistent identifier (PID) services, metadata management, and scholarly communication tools, "
        "specifically tailored for African academic and research institutions."
    )
    
    doc.add_heading('Key Features', level=2)
    features = [
        "Persistent Identifier (PID) management for DOIs, DocIDs, and Handles",
        "Integration with multiple external services (Crossref, CSTR, CORDRA, ROR, ORCID)",
        "Hierarchical commenting system with moderation capabilities",
        "Multi-provider social authentication (Google, ORCID, GitHub)",
        "Comprehensive publication metadata management",
        "RESTful API with Swagger/OpenAPI documentation",
        "Production-ready deployment with monitoring and logging"
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    add_image_placeholder(doc, "DOCiD Dashboard Screenshot", 6, 4)
    
    # 2. System Architecture
    doc.add_heading('2. System Architecture', level=1)
    
    doc.add_heading('Application Architecture', level=2)
    doc.add_paragraph(
        "The DOCiD system follows a layered architecture pattern with clear separation of concerns:"
    )
    
    arch_layers = [
        "Presentation Layer: RESTful API endpoints with JSON communication",
        "Service Layer: Business logic and external service integrations", 
        "Data Access Layer: SQLAlchemy ORM with PostgreSQL database",
        "Infrastructure Layer: Authentication, logging, caching, and monitoring"
    ]
    
    for layer in arch_layers:
        doc.add_paragraph(layer, style='List Bullet')
    
    add_image_placeholder(doc, "System Architecture Diagram", 7, 5)
    
    doc.add_heading('Technology Stack', level=2)
    
    # Create technology stack table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology'
    
    tech_stack = [
        ('Web Framework', 'Flask 3.0.3'),
        ('Database', 'PostgreSQL 12+'),
        ('ORM', 'SQLAlchemy 2.0.30'),
        ('Authentication', 'Flask-JWT-Extended'),
        ('API Documentation', 'Flasgger (Swagger/OpenAPI)'),
        ('Caching', 'Redis'),
        ('Web Server', 'Gunicorn'),
        ('Reverse Proxy', 'Nginx'),
        ('Process Management', 'Supervisor'),
        ('Migration Tool', 'Alembic')
    ]
    
    for component, technology in tech_stack:
        row_cells = table.add_row().cells
        row_cells[0].text = component
        row_cells[1].text = technology
    
    add_image_placeholder(doc, "Technology Stack Visualization", 6, 4)
    
    # 3. Installation & Setup
    doc.add_heading('3. Installation & Setup', level=1)
    
    doc.add_heading('System Requirements', level=2)
    
    # System requirements table
    req_table = doc.add_table(rows=1, cols=2)
    req_table.style = 'Table Grid'
    req_hdr = req_table.rows[0].cells
    req_hdr[0].text = 'Component'
    req_hdr[1].text = 'Requirement'
    
    requirements = [
        ('Operating System', 'Ubuntu 20.04 LTS or later'),
        ('Python', '3.9 or later'),
        ('Database', 'PostgreSQL 12 or later'),
        ('Memory', '4GB RAM minimum, 8GB recommended'),
        ('Storage', '50GB minimum for production'),
        ('Network', 'HTTPS/SSL certificate for production')
    ]
    
    for component, requirement in requirements:
        req_row = req_table.add_row().cells
        req_row[0].text = component
        req_row[1].text = requirement
    
    doc.add_heading('Development Environment Setup', level=2)
    
    doc.add_paragraph("1. Clone the repository and set up virtual environment:")
    doc.add_paragraph(
        "git clone <repository-url>\n"
        "cd DOCiD/backend\n"
        "python3 -m venv venv\n"
        "source venv/bin/activate\n"
        "pip install -r requirements.txt",
        style='Code'
    )
    
    add_image_placeholder(doc, "Development Environment Setup Terminal", 6, 3)
    
    doc.add_paragraph("2. Configure environment variables in .env file:")
    doc.add_paragraph(
        "SQLALCHEMY_DATABASE_URI=postgresql://user:pass@localhost/docid_db\n"
        "SECRET_KEY=your_secret_key\n"
        "CSTR_CLIENT_ID=your_cstr_client_id\n"
        "CROSSREF_API_KEY=your_crossref_api_key",
        style='Code'
    )
    
    doc.add_paragraph("3. Initialize database:")
    doc.add_paragraph(
        "python run.py db init\n"
        "python run.py db migrate -m 'Initial migration'\n"
        "python run.py db upgrade\n"
        "python scripts/seed_db.py",
        style='Code'
    )
    
    add_image_placeholder(doc, "Database Migration Output", 6, 3)
    
    # 4. Database Schema & Data Models
    doc.add_heading('4. Database Schema & Data Models', level=1)
    
    doc.add_heading('Core Data Models', level=2)
    
    doc.add_paragraph(
        "The DOCiD system uses a sophisticated relational database schema with the following core entities:"
    )
    
    models = [
        "UserAccount: User management with social authentication",
        "Publications: Core publication entity with metadata",
        "PublicationCreators: Author/creator information with ORCID",
        "PublicationOrganization: Institutional affiliations with ROR",
        "PublicationFunders: Funding information and grants",
        "PublicationComments: Hierarchical commenting system",
        "Reference Tables: Controlled vocabularies and lookup data"
    ]
    
    for model in models:
        doc.add_paragraph(model, style='List Bullet')
    
    add_image_placeholder(doc, "Database Entity Relationship Diagram", 8, 6)
    
    doc.add_heading('Publications Model Schema', level=2)
    
    # Publications schema table
    pub_table = doc.add_table(rows=1, cols=4)
    pub_table.style = 'Table Grid'
    pub_hdr = pub_table.rows[0].cells
    pub_hdr[0].text = 'Field'
    pub_hdr[1].text = 'Type'
    pub_hdr[2].text = 'Description'
    pub_hdr[3].text = 'Constraints'
    
    pub_fields = [
        ('id', 'Integer', 'Primary key', 'AUTO_INCREMENT'),
        ('title', 'Text', 'Publication title', 'NOT NULL'),
        ('docid', 'String(50)', 'Document identifier', 'UNIQUE'),
        ('doi', 'String(100)', 'Digital Object Identifier', 'UNIQUE'),
        ('resource_type_id', 'Integer', 'Resource type reference', 'FOREIGN KEY'),
        ('user_id', 'Integer', 'Owner user ID', 'FOREIGN KEY'),
        ('metadata', 'JSON', 'Additional metadata', 'NULL'),
        ('status', 'String(20)', 'Publication status', 'DEFAULT: draft'),
        ('created_at', 'DateTime', 'Creation timestamp', 'DEFAULT: NOW()')
    ]
    
    for field, field_type, desc, constraints in pub_fields:
        pub_row = pub_table.add_row().cells
        pub_row[0].text = field
        pub_row[1].text = field_type
        pub_row[2].text = desc
        pub_row[3].text = constraints
    
    doc.add_heading('Comments System Schema', level=2)
    
    doc.add_paragraph(
        "The commenting system implements a hierarchical structure allowing for threaded discussions:"
    )
    
    # Comments schema table
    comment_table = doc.add_table(rows=1, cols=3)
    comment_table.style = 'Table Grid'
    comment_hdr = comment_table.rows[0].cells
    comment_hdr[0].text = 'Field'
    comment_hdr[1].text = 'Type'
    comment_hdr[2].text = 'Description'
    
    comment_fields = [
        ('id', 'Integer', 'Primary key'),
        ('publication_id', 'Integer', 'Associated publication'),
        ('user_id', 'Integer', 'Comment author'),
        ('parent_comment_id', 'Integer', 'Parent comment (for replies)'),
        ('comment_text', 'Text', 'Comment content'),
        ('comment_type', 'String(50)', 'Type: general, review, question'),
        ('status', 'String(20)', 'Status: active, edited, deleted'),
        ('likes_count', 'Integer', 'Number of likes'),
        ('created_at', 'DateTime', 'Creation timestamp'),
        ('updated_at', 'DateTime', 'Last update timestamp')
    ]
    
    for field, field_type, desc in comment_fields:
        comment_row = comment_table.add_row().cells
        comment_row[0].text = field
        comment_row[1].text = field_type
        comment_row[2].text = desc
    
    add_image_placeholder(doc, "Comments Database Schema Diagram", 7, 4)
    
    # 5. API Documentation
    doc.add_heading('5. API Documentation', level=1)
    
    doc.add_heading('API Design Principles', level=2)
    
    principles = [
        "RESTful architecture with resource-oriented URLs",
        "HTTP verb semantics (GET, POST, PUT, DELETE)",
        "JSON payload communication",
        "Consistent error response format",
        "JWT-based authentication",
        "Rate limiting on sensitive endpoints",
        "Comprehensive input validation"
    ]
    
    for principle in principles:
        doc.add_paragraph(principle, style='List Bullet')
    
    add_image_placeholder(doc, "Swagger API Documentation Interface", 7, 5)
    
    doc.add_heading('Authentication Flow', level=2)
    
    doc.add_paragraph("1. User Login Request:")
    doc.add_paragraph(
        "POST /api/v1/auth/login\n"
        "Content-Type: application/json\n\n"
        "{\n"
        '  "email": "user@example.com",\n'
        '  "password": "password"\n'
        "}",
        style='Code'
    )
    
    doc.add_paragraph("2. Successful Authentication Response:")
    doc.add_paragraph(
        "{\n"
        '  "access_token": "eyJ0eXAiOiJKV1QiLCJhbGciOiJIUzI1NiJ9...",\n'
        '  "refresh_token": "eyJ0eXAiOiJKV1QiLCJhbGciOiJIUzI1NiJ9...",\n'
        '  "user": {\n'
        '    "user_id": 1,\n'
        '    "full_name": "John Doe",\n'
        '    "email": "user@example.com",\n'
        '    "role": "user"\n'
        '  }\n'
        "}",
        style='Code'
    )
    
    add_image_placeholder(doc, "Authentication Flow Diagram", 6, 4)
    
    doc.add_heading('Core API Endpoints', level=2)
    
    # API endpoints table
    api_table = doc.add_table(rows=1, cols=4)
    api_table.style = 'Table Grid'
    api_hdr = api_table.rows[0].cells
    api_hdr[0].text = 'Endpoint'
    api_hdr[1].text = 'Method'
    api_hdr[2].text = 'Description'
    api_hdr[3].text = 'Authentication'
    
    endpoints = [
        ('/api/v1/auth/login', 'POST', 'User authentication', 'None'),
        ('/api/v1/publications/publish', 'POST', 'Create publication', 'JWT'),
        ('/api/v1/publications/get-publications', 'GET', 'List publications', 'JWT'),
        ('/api/publications/{id}/comments', 'GET', 'Get comments', 'None'),
        ('/api/publications/{id}/comments', 'POST', 'Add comment', 'JWT'),
        ('/api/comments/{id}', 'PUT', 'Edit comment', 'JWT'),
        ('/api/comments/{id}', 'DELETE', 'Delete comment', 'JWT'),
        ('/api/v1/crossref/doi/{doi}', 'GET', 'Get DOI metadata', 'JWT'),
        ('/api/v1/ror/search', 'GET', 'Search organizations', 'JWT')
    ]
    
    for endpoint, method, desc, auth in endpoints:
        api_row = api_table.add_row().cells
        api_row[0].text = endpoint
        api_row[1].text = method
        api_row[2].text = desc
        api_row[3].text = auth
    
    # 6. External Service Integrations
    doc.add_heading('6. External Service Integrations', level=1)
    
    doc.add_heading('Crossref Integration', level=2)
    
    doc.add_paragraph(
        "The Crossref integration provides DOI metadata retrieval, validation, and registration capabilities. "
        "This service is essential for scholarly publication management and citation tracking."
    )
    
    crossref_features = [
        "DOI metadata retrieval and validation",
        "Publication search across Crossref database",
        "XML submission for DOI registration",
        "Bulk DOI processing capabilities",
        "Citation count and metrics retrieval"
    ]
    
    for feature in crossref_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    add_image_placeholder(doc, "Crossref Integration Architecture", 6, 4)
    
    doc.add_heading('CORDRA Integration', level=2)
    
    doc.add_paragraph(
        "CORDRA (Content Organization and Repository Digital Archive) integration enables digital object "
        "repository management and Handle identifier generation. Recent updates have enhanced identifier "
        "support for creators, funders, and organizations."
    )
    
    doc.add_heading('Enhanced Identifier Support', level=3)
    
    doc.add_paragraph(
        "The CORDRA integration now includes comprehensive identifier support with the following enhancements:"
    )
    
    cordra_enhancements = [
        "PublicationCreators: ORCID, ISNI, VIAF identifier support",
        "PublicationFunders: ROR, FundRef, ISNI identifier support", 
        "PublicationOrganizations: Future-ready for identifier fields",
        "Full resolvable URLs for all identifiers",
        "Dynamic field access for quick identifier retrieval",
        "Automatic type detection and validation"
    ]
    
    for enhancement in cordra_enhancements:
        doc.add_paragraph(enhancement, style='List Bullet')
    
    doc.add_heading('CORDRA Data Format Examples', level=3)
    
    doc.add_paragraph("Creator with ORCID Identifier:")
    doc.add_paragraph(
        "{\n"
        '  "familyName": "Kariuki",\n'
        '  "givenName": "Erastus",\n'
        '  "fullName": "Erastus Kariuki",\n'
        '  "identifier": "https://orcid.org/0000-0002-7453-6460",\n'
        '  "identifierType": "orcid",\n'
        '  "orcid": "https://orcid.org/0000-0002-7453-6460",\n'
        '  "role": "Author",\n'
        '  "parentId": "20.500.14351/..."\n'
        "}",
        style='Code'
    )
    
    doc.add_paragraph("Funder with ROR Identifier:")
    doc.add_paragraph(
        "{\n"
        '  "name": "Bill & Melinda Gates Foundation",\n'
        '  "type": "Foundation",\n'
        '  "identifier": "https://ror.org/01bj3aw27",\n'
        '  "identifierType": "ror",\n'
        '  "ror": "https://ror.org/01bj3aw27",\n'
        '  "parentId": "20.500.14351/..."\n'
        "}",
        style='Code'
    )
    
    doc.add_heading('CORDRA Push Operations', level=3)
    
    doc.add_paragraph("Manual Push Commands:")
    doc.add_paragraph(
        "# Push specific publication\n"
        "python push_to_cordra.py --publication-id 12\n\n"
        "# Push recent publications\n"
        "python push_recent_to_cordra.py\n\n"
        "# Test identifier formatting\n"
        "python test_cordra_identifiers.py",
        style='Code'
    )
    
    doc.add_paragraph(
        "The system includes automated cron jobs that push recent publications to CORDRA "
        "every minute, ensuring real-time synchronization of metadata and identifiers."
    )
    
    add_image_placeholder(doc, "CORDRA Service Configuration", 6, 3)
    add_image_placeholder(doc, "CORDRA Identifier Mapping Flow", 7, 4)
    
    doc.add_heading('ROR & ORCID Integration', level=2)
    
    doc.add_paragraph(
        "Research Organization Registry (ROR) and ORCID integrations provide institutional and researcher "
        "identification services for enhanced metadata quality."
    )
    
    add_image_placeholder(doc, "ROR/ORCID Integration Flow", 6, 4)
    
    # 7. Comments System
    doc.add_heading('7. Comments System', level=1)
    
    doc.add_heading('System Overview', level=2)
    
    doc.add_paragraph(
        "The DOCiD comments system provides a sophisticated platform for scholarly discourse and peer review. "
        "It supports hierarchical threaded discussions with moderation capabilities and engagement tracking."
    )
    
    comment_features = [
        "Hierarchical comment structure with unlimited nesting",
        "Comment types: general, review, question, suggestion",
        "Status management: active, edited, deleted, flagged",
        "Like/reaction system with count tracking",
        "User and admin moderation capabilities",
        "Edit tracking and version history",
        "Real-time statistics and analytics"
    ]
    
    for feature in comment_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    add_image_placeholder(doc, "Comments System Interface", 7, 5)
    
    doc.add_heading('Comment API Usage Examples', level=2)
    
    doc.add_paragraph("1. Retrieve Publication Comments:")
    doc.add_paragraph(
        "GET /api/publications/1/comments?include_replies=true\n\n"
        "Response:\n"
        "{\n"
        '  "publication_id": 1,\n'
        '  "total_comments": 5,\n'
        '  "comments": [\n'
        '    {\n'
        '      "id": 1,\n'
        '      "user_name": "John Doe",\n'
        '      "comment_text": "Great research work!",\n'
        '      "likes_count": 3,\n'
        '      "replies": [...]\n'
        '    }\n'
        '  ]\n'
        "}",
        style='Code'
    )
    
    doc.add_paragraph("2. Add New Comment:")
    doc.add_paragraph(
        "POST /api/publications/1/comments\n"
        "Authorization: Bearer {jwt_token}\n\n"
        "{\n"
        '  "user_id": 1,\n'
        '  "comment_text": "This is a valuable contribution.",\n'
        '  "comment_type": "review"\n'
        "}",
        style='Code'
    )
    
    add_image_placeholder(doc, "Comment Creation Interface", 6, 4)
    
    # 8. Authentication & Security
    doc.add_heading('8. Authentication & Security', level=1)
    
    doc.add_heading('Authentication Methods', level=2)
    
    auth_methods = [
        "JWT (JSON Web Tokens) for API authentication",
        "Google OAuth 2.0 integration",
        "ORCID OAuth 2.0 for researchers",
        "GitHub OAuth 2.0 integration",
        "Email/password with secure hashing",
        "Multi-factor authentication support",
        "Session management and token refresh"
    ]
    
    for method in auth_methods:
        doc.add_paragraph(method, style='List Bullet')
    
    add_image_placeholder(doc, "Authentication Methods Overview", 6, 4)
    
    doc.add_heading('Security Measures', level=2)
    
    security_measures = [
        "Password hashing with bcrypt",
        "HTTPS/SSL encryption for all communications",
        "Rate limiting on sensitive endpoints",
        "Input validation and sanitization",
        "SQL injection prevention via ORM",
        "XSS protection with output encoding",
        "CORS configuration for cross-origin security",
        "Regular security updates and monitoring"
    ]
    
    for measure in security_measures:
        doc.add_paragraph(measure, style='List Bullet')
    
    # 9. Deployment Guide
    doc.add_heading('9. Deployment Guide', level=1)
    
    doc.add_heading('Production Architecture', level=2)
    
    doc.add_paragraph(
        "The production deployment utilizes a multi-tier architecture with load balancing, "
        "SSL termination, and comprehensive monitoring."
    )
    
    add_image_placeholder(doc, "Production Deployment Architecture", 8, 6)
    
    doc.add_heading('Deployment Steps', level=2)
    
    deployment_steps = [
        "1. Server provisioning and OS configuration",
        "2. PostgreSQL database installation and configuration",
        "3. Application deployment with Gunicorn",
        "4. Nginx reverse proxy configuration",
        "5. SSL certificate installation",
        "6. Supervisor process management setup",
        "7. Background service configuration",
        "8. Monitoring and logging setup",
        "9. Backup and disaster recovery configuration",
        "10. Performance optimization and tuning"
    ]
    
    for step in deployment_steps:
        doc.add_paragraph(step)
    
    doc.add_heading('Server Configuration', level=2)
    
    doc.add_paragraph("Nginx Configuration Example:")
    doc.add_paragraph(
        "server {\n"
        "    listen 80;\n"
        "    server_name your-domain.com;\n"
        "    location / {\n"
        "        proxy_pass http://127.0.0.1:5001;\n"
        "        proxy_set_header Host $host;\n"
        "        proxy_set_header X-Real-IP $remote_addr;\n"
        "    }\n"
        "}",
        style='Code'
    )
    
    add_image_placeholder(doc, "Server Configuration Dashboard", 6, 4)
    
    # 10. Testing & Quality Assurance
    doc.add_heading('10. Testing & Quality Assurance', level=1)
    
    doc.add_heading('Testing Strategy', level=2)
    
    testing_approaches = [
        "Unit testing for individual components",
        "Integration testing for external services",
        "API endpoint testing with automated validation",
        "Database migration testing",
        "Performance testing under load",
        "Security testing and vulnerability assessment",
        "User acceptance testing for key workflows"
    ]
    
    for approach in testing_approaches:
        doc.add_paragraph(approach, style='List Bullet')
    
    doc.add_heading('Test Execution', level=2)
    
    doc.add_paragraph("Run the comprehensive test suite:")
    doc.add_paragraph(
        "# API endpoint testing\n"
        "python test_comments_api.py\n"
        "python test_comments_fetch.py\n\n"
        "# Service integration testing\n"
        "python test_single_cstr.py\n\n"
        "# Error diagnosis\n"
        "python diagnose_comments_error.py",
        style='Code'
    )
    
    add_image_placeholder(doc, "Test Results Dashboard", 7, 4)
    
    # 11. Monitoring & Maintenance
    doc.add_heading('11. Monitoring & Maintenance', level=1)
    
    doc.add_heading('Monitoring Setup', level=2)
    
    monitoring_components = [
        "Application performance monitoring (APM)",
        "Database performance tracking",
        "Server resource monitoring (CPU, memory, disk)",
        "API endpoint response time tracking",
        "Error rate monitoring and alerting",
        "External service integration monitoring",
        "Log aggregation and analysis",
        "Uptime monitoring and notifications"
    ]
    
    for component in monitoring_components:
        doc.add_paragraph(component, style='List Bullet')
    
    add_image_placeholder(doc, "Monitoring Dashboard", 7, 5)
    
    doc.add_heading('Maintenance Procedures', level=2)
    
    doc.add_paragraph("Regular maintenance tasks include:")
    
    maintenance_tasks = [
        "Database backup and recovery testing",
        "Security updates and patch management",
        "Performance optimization and query tuning",
        "Log rotation and cleanup",
        "SSL certificate renewal",
        "External service credential rotation",
        "Capacity planning and scaling assessment"
    ]
    
    for task in maintenance_tasks:
        doc.add_paragraph(task, style='List Bullet')
    
    # 12. Appendices
    doc.add_heading('12. Appendices', level=1)
    
    doc.add_heading('Appendix A: Environment Variables', level=2)
    
    # Environment variables table
    env_table = doc.add_table(rows=1, cols=3)
    env_table.style = 'Table Grid'
    env_hdr = env_table.rows[0].cells
    env_hdr[0].text = 'Variable'
    env_hdr[1].text = 'Description'
    env_hdr[2].text = 'Required'
    
    env_vars = [
        ('SQLALCHEMY_DATABASE_URI', 'Database connection string', 'Yes'),
        ('SECRET_KEY', 'Application secret key', 'Yes'),
        ('JWT_SECRET_KEY', 'JWT token secret', 'Yes'),
        ('CSTR_CLIENT_ID', 'CSTR service client ID', 'Yes'),
        ('CSTR_SECRET', 'CSTR service secret', 'Yes'),
        ('CROSSREF_API_KEY', 'Crossref API key', 'Yes'),
        ('MAIL_SERVER', 'SMTP server address', 'No'),
        ('GOOGLE_CLIENT_ID', 'Google OAuth client ID', 'No'),
        ('ORCID_CLIENT_ID', 'ORCID OAuth client ID', 'No')
    ]
    
    for var, desc, required in env_vars:
        env_row = env_table.add_row().cells
        env_row[0].text = var
        env_row[1].text = desc
        env_row[2].text = required
    
    doc.add_heading('Appendix B: Error Codes', level=2)
    
    # Error codes table
    error_table = doc.add_table(rows=1, cols=2)
    error_table.style = 'Table Grid'
    error_hdr = error_table.rows[0].cells
    error_hdr[0].text = 'Error Code'
    error_hdr[1].text = 'Description'
    
    error_codes = [
        ('INVALID_CREDENTIALS', 'Login credentials incorrect'),
        ('USER_NOT_FOUND', 'User account does not exist'),
        ('PUBLICATION_NOT_FOUND', 'Publication does not exist'),
        ('COMMENT_NOT_FOUND', 'Comment does not exist'),
        ('UNAUTHORIZED_ACCESS', 'Insufficient permissions'),
        ('VALIDATION_ERROR', 'Input validation failed'),
        ('EXTERNAL_SERVICE_ERROR', 'External API integration error'),
        ('DATABASE_ERROR', 'Database operation failed')
    ]
    
    for code, desc in error_codes:
        error_row = error_table.add_row().cells
        error_row[0].text = code
        error_row[1].text = desc
    
    doc.add_heading('Appendix C: Command Reference', level=2)
    
    doc.add_paragraph("Development Commands:")
    doc.add_paragraph(
        "# Database operations\n"
        "python run.py db migrate -m 'Migration message'\n"
        "python run.py db upgrade\n"
        "python run.py db downgrade\n\n"
        "# Service management\n"
        "./setup_cstr_service.sh\n"
        "python push_to_cordra.py\n\n"
        "# Testing\n"
        "python test_comments_api.py\n"
        "python diagnose_comments_error.py",
        style='Code'
    )
    
    add_image_placeholder(doc, "Command Line Interface Examples", 6, 3)
    
    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("End of Documentation")
    footer_run.font.size = Pt(10)
    footer_run.italic = True
    
    # Save document
    filename = f"/Users/ekariz/Projects/AMBAND/DOCiD/backend/DOCiD_API_Documentation_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
    doc.save(filename)
    
    print(f"✅ Documentation created successfully: {filename}")
    print(f"📊 Document contains {len(doc.paragraphs)} paragraphs")
    print(f"📷 Image placeholders added for screenshots")
    print(f"📋 Comprehensive coverage of all system components")
    
    return filename

if __name__ == "__main__":
    create_docx_documentation()